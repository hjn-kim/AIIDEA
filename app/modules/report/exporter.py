import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.modules.report.model import Report
from app.file_utils import get_download_path


def _sanitize(name: str, default: str = "보고서") -> str:
    name = (name or "").strip()
    if not name:
        return default
    return re.sub(r'[\\/:*?"<>|]+', "_", name)[:120]


def save_report_docx(r: Report, toc: str, sections: dict, sources: list) -> str:
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ── 표지 ──────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(r.title)
    run.bold = True
    run.font.size = Pt(32)
    run.font.name = "맑은 고딕"
    run.font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{r.author}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = []
    if r.date:
        lines.append(f"작성일: {r.date}")
    run = p.add_run("\n".join(lines))
    run.font.size = Pt(10)
    run.font.name = "맑은 고딕"
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # ── 목차 ──────────────────────────────────────────
    h = doc.add_heading("목차", level=1)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
    _chapter_re = re.compile(r"^\d+\s*[\.\)]?\s+\S")
    _sub_re = re.compile(r"^\d+\.\d+")
    prev_was_chapter = False
    for line in (toc or "").splitlines():
        s = line.strip()
        if not s:
            continue
        is_chapter = bool(_chapter_re.match(s)) and not _sub_re.match(s)
        if is_chapter and prev_was_chapter:
            doc.add_paragraph()
        p = doc.add_paragraph()
        if _sub_re.match(s):
            p.paragraph_format.left_indent = Pt(20)
        p.add_run(s)
        prev_was_chapter = is_chapter

    doc.add_page_break()

    # ── 본문 ──────────────────────────────────────────
    _body_chapter_re = re.compile(r"^\d+\s*[\.\)]?\s+\S")
    _body_sub_re = re.compile(r"^\d+\.\d+")
    for chapter_title, content in sections.items():
        is_major = bool(_body_chapter_re.match(chapter_title.strip())) and not _body_sub_re.match(chapter_title.strip())
        if is_major:
            doc.add_page_break()
        ch = doc.add_heading(chapter_title, level=2)
        if ch.runs:
            ch.runs[0].font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
        if content.strip():
            for line in content.splitlines():
                if line.strip():
                    doc.add_paragraph(line)
        doc.add_paragraph()

    # ── 참고문헌 ──────────────────────────────────────
    if sources:
        h = doc.add_heading("참고문헌", level=1)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3C, 0x78)
        for src in sources:
            doc.add_paragraph(src)

    file_name = f"{_sanitize(r.title)}.docx"
    path = get_download_path(file_name)
    doc.save(path)
    return path
