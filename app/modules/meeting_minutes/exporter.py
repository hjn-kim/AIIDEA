import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.modules.meeting_minutes.model import MeetingMinutes, parse_ai_sections
from app.file_utils import get_download_path


NAVY = RGBColor(0x1A, 0x3A, 0x6B)
GRAY = RGBColor(80, 80, 80)


def _sanitize(name: str, default: str = "회의록") -> str:
    name = (name or "").strip()
    if not name:
        return default
    return re.sub(r'[\\/:*?"<>|]+', "_", name)[:120]


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _font(run, size_pt=10, bold=False, color: RGBColor = None, name="맑은 고딕"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _para(doc, text="", size_pt=10, bold=False, color: RGBColor = None,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=4, indent_cm=0.0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if text:
        _font(p.add_run(text), size_pt=size_pt, bold=bold, color=color)
    return p


def export_docx(m: MeetingMinutes, ai_text: str) -> str:
    doc = Document()

    # 페이지 여백
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

    # 기본 폰트
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)

    # ── 표지 제목 ──────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph()

    _para(doc, "회  의  록", size_pt=30, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

    _para(doc, m.title or "", size_pt=14, color=GRAY,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

    # ── 기본 정보 표 ───────────────────────────────────
    info_rows = [
        ("주최측", m.host),
        ("회의 목적", m.purpose),
        ("일시", m.meeting_datetime),
        ("장소", m.location),
        ("참석자", m.attendees),
    ]
    info_rows = [(k, v) for k, v in info_rows if v and v.strip()]

    if info_rows:
        tbl = doc.add_table(rows=len(info_rows), cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, (k, v) in enumerate(info_rows):
            kc = tbl.rows[i].cells[0]
            vc = tbl.rows[i].cells[1]
            _set_cell_bg(kc, "DCE6F1")
            kp = kc.paragraphs[0]
            kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(kp.add_run(k), size_pt=10, bold=True, color=NAVY)
            _font(vc.paragraphs[0].add_run(v), size_pt=10)
            kc.width = Cm(3.5)
            vc.width = Cm(11.5)
            tbl.rows[i].height = Cm(1.0)
            tbl.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    doc.add_page_break()

    # ── 본문: AI 생성 섹션 ─────────────────────────────
    sections = parse_ai_sections(ai_text)

    if not sections:
        # 파싱 실패 시 전문 그대로 출력
        for line in ai_text.splitlines():
            s = line.strip()
            if s:
                _para(doc, s, size_pt=10.5, after=4)
    else:
        for heading, body in sections:
            # 섹션 제목
            _para(doc, heading, size_pt=13, bold=True, color=NAVY,
                  before=10, after=4)
            # 구분선
            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(6)
            _font(sep.add_run("─" * 55), size_pt=7, color=GRAY)

            # 본문 내용
            for line in body.splitlines():
                s = line.strip()
                if not s:
                    continue
                is_bullet = s.startswith("-") or s.startswith("•")
                if is_bullet:
                    text = "•  " + s.lstrip("-•* ").strip()
                    _para(doc, text, size_pt=10.5, after=3, indent_cm=0.7)
                else:
                    _para(doc, s, size_pt=10.5, after=3)

            doc.add_paragraph()

    file_name = f"{_sanitize(m.title)}.docx"
    path = get_download_path(file_name)
    doc.save(path)
    return path
