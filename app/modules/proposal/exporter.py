import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.modules.proposal.model import Proposal
from app.file_utils import get_download_path


# ── 스타일 헬퍼 ───────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """셀 배경색 설정"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _font(run, name="맑은 고딕", size_pt=None, bold=False, color_rgb=None):
    """폰트 적용 (한글 폰트 포함)"""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def _fmt(para, before=0, after=6, line=1.2):
    """단락 간격/줄간격 설정"""
    f = para.paragraph_format
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing = line


def _is_table_line(line: str) -> bool:
    return line.strip().startswith("|")


def _is_separator_row(cells: list) -> bool:
    return all(re.match(r"^[-: ]+$", c) for c in cells if c)


def _parse_md_table(lines: list) -> list:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if _is_separator_row(cells):
            continue
        rows.append(cells)
    return rows


def _add_md_table(doc, rows_data: list):
    if not rows_data:
        return
    n_cols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=n_cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows_data):
        for ci in range(n_cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = tbl.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if ri == 0:
                _font(run, size_pt=10, bold=True, color_rgb=(0x1A, 0x3A, 0x6B))
                _set_cell_bg(cell, "DCE6F1")
            else:
                _font(run, size_pt=10)
    doc.add_paragraph()


# ── 메인 내보내기 ─────────────────────────────────────────────────

def export_docx(p: Proposal) -> str:
    doc = Document()

    # 페이지 여백
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)

    BLUE  = (0x1A, 0x3A, 0x6B)
    LGRAY = (0xBB, 0xBB, 0xBB)

    # ── 표지 ────────────────────────────────────────────────────
    for _ in range(7):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(title_p.add_run(p.title or "제  안  서"), size_pt=28, bold=True, color_rgb=BLUE)
    _fmt(title_p, before=0, after=18)

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(line_p.add_run("─" * 28), size_pt=10, color_rgb=BLUE)
    _fmt(line_p, before=0, after=24)

    # 기본 정보 표
    info_rows = [(k, v) for k, v in [
        ("제안 일자", p.date),
        ("제안사",   p.from_company),
        ("수신사",   p.to_company),
    ] if v and v.strip()]

    if info_rows:
        tbl = doc.add_table(rows=len(info_rows), cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, (k, v) in enumerate(info_rows):
            kc = tbl.rows[i].cells[0]
            vc = tbl.rows[i].cells[1]
            _set_cell_bg(kc, "DCE6F1")
            kp = kc.paragraphs[0]
            kp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(kp.add_run(k), size_pt=10, bold=True, color_rgb=BLUE)
            _font(vc.paragraphs[0].add_run(v), size_pt=10)
            kc.width = Cm(4)
            vc.width = Cm(11)
            tbl.rows[i].height = Cm(1.2)
            tbl.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    doc.add_page_break()

    # ── 목차 ────────────────────────────────────────────────────
    toc_title = doc.add_paragraph()
    _font(toc_title.add_run("목  차"), size_pt=14, bold=True, color_rgb=BLUE)
    _fmt(toc_title, before=0, after=10)

    toc_sep = doc.add_paragraph()
    _font(toc_sep.add_run("─" * 60), size_pt=7, color_rgb=LGRAY)
    _fmt(toc_sep, before=0, after=10)

    for i, section_name in enumerate(p.sections.keys(), 1):
        toc_item = doc.add_paragraph()
        toc_item.paragraph_format.left_indent = Cm(0.5)
        _font(toc_item.add_run(f"{i}.  {section_name}"), size_pt=10.5)
        _fmt(toc_item, before=0, after=5)

    doc.add_page_break()

    # ── 섹션별 본문 ──────────────────────────────────────────────
    for idx, (sec_name, content) in enumerate(p.sections.items()):
        # 섹션 번호 + 제목
        h = doc.add_paragraph()
        _font(h.add_run(f"{idx + 1}.  {sec_name}"), size_pt=13, bold=True, color_rgb=BLUE)
        _fmt(h, before=8 if idx > 0 else 0, after=3)

        # 구분선
        sep = doc.add_paragraph()
        _font(sep.add_run("─" * 60), size_pt=7, color_rgb=LGRAY)
        _fmt(sep, before=0, after=8)

        # 본문 내용
        if content.strip():
            raw_lines = content.splitlines()
            i = 0
            while i < len(raw_lines):
                line = raw_lines[i]
                stripped = line.strip()

                # 마크다운 표 블록 수집
                if _is_table_line(stripped):
                    table_lines = []
                    while i < len(raw_lines) and _is_table_line(raw_lines[i].strip()):
                        table_lines.append(raw_lines[i].strip())
                        i += 1
                    _add_md_table(doc, _parse_md_table(table_lines))
                    continue

                if not stripped:
                    sp = doc.add_paragraph()
                    _fmt(sp, before=0, after=2)
                    i += 1
                    continue

                is_bullet = stripped[0] in "-•*"
                bp = doc.add_paragraph()
                if is_bullet:
                    bp.paragraph_format.left_indent = Cm(0.8)
                    bp.paragraph_format.first_line_indent = Cm(-0.4)
                    text = "•  " + stripped.lstrip("-•* ").strip()
                else:
                    text = stripped
                _font(bp.add_run(text), size_pt=10.5)
                _fmt(bp, before=0, after=4, line=1.35)
                i += 1

        # 섹션 간 빈 단락
        doc.add_paragraph()

    # 파일 저장
    safe = (p.title or "제안서").strip()
    for ch in r'/\:*?"<>|':
        safe = safe.replace(ch, "_")
    path = get_download_path(f"{safe}.docx")
    doc.save(path)
    return path
