# -*- coding: utf-8 -*-
"""
보고서 생성 엔진: 목차 생성, 목차 수정, 본문 생성, DOCX 저장
"""
import os
import re
import time
from google import genai

MODEL_MAP = {"Flash": "gemini-2.0-flash", "Pro": "gemini-2.0-pro"}


def sanitize_filename(name: str, default: str = "report") -> str:
    name = (name or "").strip()
    if not name:
        return default
    name = re.sub(r'[\\/:*?"<>|]+', "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name[:120]


def get_download_path(file_name: str = "report.docx") -> str:
    if os.name == "nt":
        download_path = os.path.join(os.environ.get("USERPROFILE", ""), "Downloads")
    else:
        download_path = os.path.join(os.environ.get("HOME", ""), "Downloads")
    if not download_path or not os.path.isdir(download_path):
        download_path = os.getcwd()
    os.makedirs(download_path, exist_ok=True)
    return os.path.join(download_path, file_name)


def _make_prompt_toc(topic: str, purpose: str, requirements: str) -> str:
    return f"""
너는 컨설팅 보고서 작성 전문가다.
아래 입력을 바탕으로 '보고서 목차(TOC)'를 한국어로 작성하라.

[입력]
- 보고서 제목: {topic}
- 보고서 목적: {purpose}
- 조건/요구사항: {requirements}

[출력 요구사항]
- 출력은 번호가 매겨진 목차 항목만 포함해야 한다.
- 제목, 설명, 안내문, 라벨을 절대 출력하지 말 것.
- "목차", "개정 목차", "개정목차"라는 문자열을 어떤 형태로도 출력하지 말 것.
- #, ##, ###, *, -, • 등 마크다운 기호를 절대 사용하지 말 것.
- 번호 체계는 1, 1.1, 1.2 형태로 작성
- 최소 5개 장(Chapter) 이상
- 장(1,2,3...)과 소절(1.1, 1.2...)로 구성
- 불필요한 설명문 없이 목차만 출력
- 도입부에 목차라는 표시 표현 생성 금지
- # 또는 * 절대 사용금지
""".strip()


def _make_prompt_revise(topic: str, purpose: str, requirements: str, toc: str, feedback: str) -> str:
    return f"""
너는 컨설팅 보고서 편집자다.
아래 기존 목차와 사용자 수정 지시를 반영하여 '개정 목차'를 작성하라.

[입력]
- 보고서 제목: {topic}
- 보고서 목적: {purpose}
- 조건/요구사항: {requirements}

[기존 목차]
{toc}

[목차 수정]
{feedback}

[출력 요구사항]
- 출력은 번호가 매겨진 목차 항목만 포함해야 한다.
- 제목, 설명, 안내문, 라벨을 절대 출력하지 말 것.
- "목차", "개정 목차", "개정목차"라는 문자열을 어떤 형태로도 출력하지 말 것.
- #, ##, ###, *, -, • 등 마크다운 기호를 절대 사용하지 말 것.
- 번호 체계는 1, 1.1, 1.2 형태로 작성
- 불필요한 설명문 없이 목차만 출력
- 도입부에 개정목차라는 표시 표현 생성 금지
- # 또는 * 기호 절대 사용하지 말것
""".strip()


def generate_toc(
    api_key: str,
    topic: str,
    purpose: str,
    requirements: str,
    model: str = "Flash",
) -> str:
    """목차 1차 생성. 반환: 목차 텍스트."""
    model_name = MODEL_MAP.get(model, "gemini-2.0-flash")
    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(
        model=model_name,
        contents=_make_prompt_toc(topic, purpose, requirements),
    )
    return (getattr(resp, "text", None) or "").strip()


def revise_toc(
    api_key: str,
    topic: str,
    purpose: str,
    requirements: str,
    current_toc: str,
    feedback: str,
    model: str = "Flash",
) -> str:
    """목차 수정 반영. 반환: 개정 목차 텍스트."""
    model_name = MODEL_MAP.get(model, "gemini-2.0-flash")
    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(
        model=model_name,
        contents=_make_prompt_revise(topic, purpose, requirements, current_toc, feedback),
    )
    return (getattr(resp, "text", None) or "").strip()


def make_prompt_for_report(
    chapter_title: str,
    topic: str,
    purpose: str,
    requirements: str,
    summary_context: str = "",
    style: str = "줄글형",
) -> str:
    return f"""
너는 대형 컨설팅펌 수준의 전문 보고서를 작성하는 최고 수준의 분석가이다.
아래 입력 정보를 바탕으로, 해당 소제목에 대해서
심층적이고 장문(긴 분량)의 고품질 보고서를 작성하라.

[입력]
- 해당 소제목: {chapter_title}
- 보고서 제목: {topic}
- 보고서 목적: {purpose}
- 보고서 요구사항: {requirements}
- 서술방식: {style}
- 이전 내용 요약: {summary_context}

[출력 요구사항]
1. REPORT
- 서술방식이 '줄글형'이면 본문을 문단 형태의 줄글로 서술하고, 'Bullet 항목형'이면 본문을 - 기호로 시작하는 불릿 항목 나열형으로 서술하라.
- "해당 소제목"에 정확히 대응하는 심층적인 본문을 작성하라.
- 단순 설명을 지양하고, 원인·구조·배경·영향·시사점의 관점에서 심층적이고 전문적인 분석을 수행하라.
- 해당 주제에 대해 객관적 사실에 기반한 정확한 분석을 제공하라.
- 관련 데이터(통계, 조사 결과, 시장 동향)와 사례(기존 연구, 산업 사례, 국가·지역 비교)를 적극 활용하여 논리를 강화하라.
- 이전 내용 요약을 참고하여 논리적 흐름이 자연스럽게 이어지도록 작성하라.
- 검증된 사실 범위 내에서 가능한 한 많은 유의미한 정보를 포함하라.
- 내용은 보고서 목적에 직접적으로 부합하도록 구성하라.
- 장 제목과 소제목은 "1. 서론", "1.1 정의"와 같은 형태로만 작성하며, ##, ### 등 마크다운 헤더 기호는 절대 사용하지 마라.
- 문자 * 및 강조를 위한 모든 기호는 절대 사용하지 마라.
- 모든 본문은 빈 줄 없이 작성하라.
- 문단은 필요한 경우에만 구분할 수 있으며, 문단이 바뀔 때에는 단일 줄바꿈만 허용한다.
- 문단 내부의 일반 서술 문장 사이에는 줄바꿈을 사용하지 마라.
- 나열이 필요한 경우에만 줄바꿈을 허용하며,
  각 항목은 반드시 -기호로 시작하고,
  항목 간에는 단일 줄바꿈만 사용하라.
- 연속된 줄바꿈(빈 줄)은 절대 사용하지 마라.
- 소제목 바로 아래에도 빈 줄을 두지 마라.
- 검증되지 않은 사실, 확인되지 않은 통계, 존재 여부가 불확실한 연구·기관·보고서·법령은 절대 작성하지 마라.
- 실제로 존재하는지 확신할 수 없는 수치·사례·정책명은 사용하지 마라.
- 특정 수치나 연구 결과를 제시하는 경우, 해당 수치가 일반적으로 널리 알려진 공신력 자료인지 확신할 수 있을 때에만 사용하라.
- 확신할 수 없는 경우에는 수치를 생성하지 말고 “구체적 통계는 별도 확인 필요”라고 명시하라.
- 분석을 위해 일반적 개념 설명이 필요한 경우에도, 실존 연구나 기관명을 임의로 생성하지 마라.
- 확실한 출처를 특정할 수 없는 경우와 불확실한 정보는 아예 작성하지 마라.

2. SUMMARY
- 해당 소제목의 핵심 내용을 정확히 2문장으로 요약하라.
- 분석 결과와 시사점이 포함되도록 하라.

3. SOURCES
- 해당 장의 본문, 하위 질문, 분석 및 답변을 작성하는 과정에서
  사실 확인, 수치 인용, 분석 틀 설정, 판단 근거로 실제 사용된 모든 외부 출처를
  MLA 형식으로 기재하라.
- 직접 인용 여부와 무관하게 분석에 활용되었으면 포함하라.
- 중복 출처는 1회만 기재하라.
- SOURCES에는 실제로 존재하는 공신력 자료만 기재하라.
- 존재 여부를 확신할 수 없는 논문, 보고서, 기관, 웹사이트는 절대 생성하지 마라.
- 특정 통계나 연구를 본문에 사용하지 않았다면 SOURCES에 포함하지 마라.
- MLA 형식으로 작성하되, 허위 DOI, 허위 URL, 허위 발행연도는 절대 생성하지 마라.
- 실제 출처를 특정할 수 없는 경우, 기재 하지 마라.
- 확실한 출처를 특정할 수 없는 경우와 불확실한 정보는 아예 작성하지 마라.

[출력 형식]
반드시 아래 3개 블록만으로 출력하라. 블록의 순서, 대괄호 태그 표기, 시작/종료 태그는 그대로 유지하라.

[REPORT]
(해당 소제목 보고서 본문)
[/REPORT]

[SUMMARY]
(해당 소제목 핵심 요약 2문장)
[/SUMMARY]

[SOURCES]
(MLA 형식 출처 목록)
[/SOURCES]
""".strip()


def extract_chapters(final_toc: str) -> list:
    chapters = []
    pattern = re.compile(r"^\s*(\d+(?:\.\d+)*)\s*[\.\)]?\s+.+?:?\s*$")
    for line in (final_toc or "").splitlines():
        s = line.strip().replace("*", "")
        if pattern.match(s):
            chapters.append(s.rstrip(":").strip())
    return chapters


def parse_response_blocks(text: str, chapter_title: str = "") -> tuple:
    t = (text or "").strip()
    if not t:
        return f"{chapter_title} 오류!", "", ""

    def between_pair(txt: str, a: str, b: str) -> str:
        if a not in txt or b not in txt:
            return ""
        return txt.split(a, 1)[1].split(b, 1)[0].strip()

    report = between_pair(t, "[REPORT]", "[/REPORT]")
    summary = between_pair(t, "[SUMMARY]", "[/SUMMARY]")
    sources = between_pair(t, "[SOURCES]", "[/SOURCES]")
    if report:
        return report, summary, sources
    if "[REPORT]" in t and "[SUMMARY]" in t and "[SOURCES]" in t:
        try:
            after_report = t.split("[REPORT]", 1)[1]
            report2 = after_report.split("[SUMMARY]", 1)[0].strip()
            after_summary = after_report.split("[SUMMARY]", 1)[1]
            summary2 = after_summary.split("[SOURCES]", 1)[0].strip()
            sources2 = after_summary.split("[SOURCES]", 1)[1].strip()
            if report2:
                return report2, summary2, sources2
        except Exception:
            pass
    return f"{chapter_title} 오류!", "", ""


def add_multiline_text(doc, text: str) -> None:
    for line in (text or "").splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)


def build_docx(
    topic: str,
    purpose: str,
    requirements: str,
    final_toc: str,
    api_key: str,
    style: str = "줄글형",
    model: str = "Flash",
    summary_context: str = "",
    dedup_sources: bool = True,
    sleep_sec: float = 2.0,
) -> str:
    """본문 생성 후 DOCX 저장. 반환: 저장된 파일 경로."""
    from docx import Document

    model_name = MODEL_MAP.get(model, "gemini-2.0-flash")
    client = genai.Client(api_key=api_key)
    chapters = extract_chapters(final_toc)

    doc = Document()
    doc.add_heading(topic, level=1)
    doc.add_heading("TOC", level=1)
    add_multiline_text(doc, final_toc)
    doc.add_heading("Report", level=1)

    updated_summary_context = summary_context or ""
    seen_sources = set()
    sources_out = []

    for chapter_title in chapters:
        prompt = make_prompt_for_report(
            chapter_title=chapter_title,
            topic=topic,
            purpose=purpose,
            requirements=requirements,
            summary_context=updated_summary_context,
            style=style,
        )
        response = client.models.generate_content(model=model_name, contents=prompt)
        time.sleep(sleep_sec)
        raw_text = getattr(response, "text", "") or ""
        r, s, src = parse_response_blocks(raw_text, chapter_title)

        doc.add_heading(chapter_title, level=2)
        add_multiline_text(doc, r or f"{chapter_title} 오류!")
        if s:
            updated_summary_context += f"- {chapter_title}: {s}\n"
        else:
            updated_summary_context += f"- {chapter_title}: (요약 없음)\n"

        if src:
            if dedup_sources:
                for line in src.splitlines():
                    item = line.strip()
                    if not item:
                        continue
                    if item not in seen_sources:
                        seen_sources.add(item)
                        sources_out.append(item)
            else:
                sources_out.append(src.strip())

    doc.add_heading("Sources", level=1)
    add_multiline_text(doc, "\n".join(sources_out).strip())
    file_name = f"{sanitize_filename(topic)}.docx"
    file_path = get_download_path(file_name)
    doc.save(file_path)
    return file_path
